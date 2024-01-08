import os

import pandas as pd
from docx import Document
from utils.pendulum_helper import parse_date, pendulum_to_string
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



def generate_document_from_ssim(country, ssim_file, airline_name, contact_person, handler): # TODO: add type hints?
    
    """
    Generates a landing permit document in .docx format for a specific country starting from a SSIM file.

    This function creates a Word document requesting a landing permit. It includes 
    details about the flights operated by an airline to the specified country within a certain date range. 
    The document contains a heading, an introductory paragraph, a table of flight series, 
    and a closing signature.

    Args:
        country (str): The country for which the landing permit is requested.
        ssim_file (SSIM_File): An object representing the SSIM file containing flight data.
        airline_name (str): The name of the airline requesting the permit.
        contact_person (str): The name of the contact person for the airline.
        handler (FlightSeriesHandler): A handler object to process flight series data.

    Returns:
        None: The function creates and saves a .docx file but does not return any value.
    """
    
    start_date = ssim_file.start_date
    end_date = ssim_file.end_date

    # Initialize python docx document
    doc = Document()

    # Add the static content
    doc.add_heading('Request for Landing Permit', 0)

    # Add a first paragraph with sample content
    doc.add_paragraph(
        ("Dear Sir/Madam,\n\n"
         "I am writing to request a landing permit for {country} "
         "for the period {start_date} to {end_date}.\n\n"
         "{airline_name} is a scheduled airline operating flights to {country}.\n\n"
         "We are planning to operate the following flights to {country}:").format(
             country=country, 
             start_date=start_date.strftime('%d %B %Y'), 
             end_date=end_date.strftime('%d %B %Y'), 
             airline_name=airline_name)
    )

    # Filter flight series by country using the handler
    flight_series = handler.filter_by_country(country)

    # Convert filtered flight series to DataFrame (assuming you have a method in FlightSeries object to return its data as a dict)
    data_dicts = [fs.to_dict() for fs in flight_series]
    df = pd.DataFrame(data_dicts)

    # Add a table with the flight series
    t = doc.add_table(df.shape[0] + 1, df.shape[1])
    
    # Add column headers
    for j, column in enumerate(df.columns):
        t.cell(0, j).text = column
        
    # Populate rest of the table with the flight series data
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            t.cell(i + 1, j).text = str(df.values[i, j])

    # Add a line break after the table
    doc.add_paragraph("\n")

    # Add signature paragraph
    doc.add_paragraph(f'Sincerely, {contact_person}')

    # Ensure directory structure exists before saving
    base_directory = os.path.join(os.getcwd(), "permits_output")
    country_directory = os.path.join(base_directory, country)
    
    if not os.path.exists(country_directory):
        os.makedirs(country_directory)
        
    doc.save(os.path.join(country_directory, f'landing_permit_{country}.docx'))


# def generate_document_from_flight_series(country, list_of_flight_series, airline_name, contact_person, handler): # TODO: add type hints?
    
#     """
#     Generates a landing permit document in .docx format for a specific country starting from a list of flight series.

#     This function creates a Word document requesting a landing permit. It includes 
#     details about the flights operated by an airline to the specified country within a certain date range. 
#     The document contains a heading, an introductory paragraph, a table of flight series, 
#     and a closing signature.

#     Args:
#         country (str): The country for which the landing permit is requested.
#         list_of_flight_series (list of FlightSeries): An object representing the SSIM file containing flight data.
#         airline_name (str): The name of the airline requesting the permit.
#         contact_person (str): The name of the contact person for the airline.
#         handler (FlightSeriesHandler): A handler object to process flight series data.

#     Returns:
#         None: The function creates and saves a .docx file but does not return any value.
#     """

#     # Determine start and end dates from flight series
#     start_date = pendulum_to_string(min([parse_date(fs.effective_date) for fs in list_of_flight_series]))
#     end_date = pendulum_to_string(max([parse_date(fs.discontinued_date) for fs in list_of_flight_series]))
    
#     # Re-initialize the handler's flight series collection
#     handler.flight_series_collection = []
    
#     # Add flight series to handler
#     for fs in list_of_flight_series:
#         handler.add_flight_series(fs.to_dict())

#     # Initialize python docx document
#     doc = Document()

#     # Add the static content
#     doc.add_heading('Request for Landing Permit', 0)

#     # Add a first paragraph with sample content
#     doc.add_paragraph(
#         ("Dear Sir/Madam,\n\n"
#          "I am writing on behalf of {airline_name} to request a landing permit to perform scheduled oprations to {country} "
#          "for the period {start_date} to {end_date}.\n\n"
#          "{airline_name} is a scheduled airline operating flights to {country}.\n\n"
#          "We are planning to operate the following flights to {country}:").format(
#              country=country, 
#              start_date=start_date, 
#              end_date=end_date, 
#              airline_name=airline_name)
#     )

#     # Filter flight series by country using the handler
#     flight_series = handler.filter_by_country(country)

#     # Convert filtered flight series to DataFrame (assuming you have a method in FlightSeries object to return its data as a dict)
#     data_dicts = [fs.to_dict() for fs in flight_series]
#     df = pd.DataFrame(data_dicts)

#     # Add a table with the flight series
#     t = doc.add_table(df.shape[0] + 1, df.shape[1])
    
#     # Add column headers
#     for j, column in enumerate(df.columns):
#         t.cell(0, j).text = column
        
#     # Populate rest of the table with the flight series data
#     for i in range(df.shape[0]):
#         for j in range(df.shape[1]):
#             t.cell(i + 1, j).text = str(df.values[i, j])

#     # Add a line break after the table
#     doc.add_paragraph("\n")

#     # Add signature paragraph
#     doc.add_paragraph(f'Sincerely, {contact_person}')

#     # Ensure directory structure exists before saving
#     base_directory = os.path.join(os.getcwd(), "permits_output")
#     country_directory = os.path.join(base_directory, country)
    
#     if not os.path.exists(country_directory):
#         os.makedirs(country_directory)
        
#     doc.save(os.path.join(country_directory, f'landing_permit_{country}.docx'))

from docx import Document
import os
import pandas as pd
import pendulum

def generate_document_from_flight_series(country, list_of_flight_series, airline_name, contact_person, handler):
    # Initialize start and end dates
    start_date, end_date = None, None

    # Re-initialize the handler's flight series collection and determine start/end dates
    handler.flight_series_collection = []

    # Add flight series to handler and determine start/end dates
    for fs in list_of_flight_series:
        handler.add_flight_series(fs.to_dict())  # Assuming batch add is not possible

        eff_date = parse_date(fs.effective_date)
        dis_date = parse_date(fs.discontinued_date)

        start_date = min(start_date, eff_date) if start_date else eff_date
        end_date = max(end_date, dis_date) if end_date else dis_date

    # Convert to string format
    start_date_str = pendulum_to_string(start_date)
    end_date_str = pendulum_to_string(end_date)

    doc = Document()
    doc.add_heading('Request for Landing Permit', 0)
    doc.add_paragraph((
        f"Dear Sir/Madam,\n\n"
        f"I am writing on behalf of {airline_name} to request a landing permit to perform scheduled operations to {country} "
        f"for the period {start_date_str} to {end_date_str}.\n\n"
        f"{airline_name} is a scheduled airline operating flights to {country}.\n\n"
        f"We are planning to operate the following flights to {country}:"
    ))

    # Filter flight series by country using the handler
    filtered_flight_series = handler.filter_by_country(country)

    # Create DataFrame for table from filtered flight series
    filtered_data_dicts = [fs.to_dict() for fs in filtered_flight_series]
    df = pd.DataFrame(filtered_data_dicts)

    # Add a table with the filtered flight series
    add_table_from_dataframe(doc, df)

    doc.add_paragraph("\nSincerely, {contact_person}")

    # Ensure directory structure exists before saving
    country_directory = os.path.join(os.getcwd(), "permits_output", country)
    os.makedirs(country_directory, exist_ok=True)
    
    doc.save(os.path.join(country_directory, f'landing_permit_{country}.docx'))



from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_cell_font(cell, font_name, font_size):
    """
    Set the font name and size for a cell.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

def set_column_widths(table, df):
    """
    Set the column widths to match the longest word in each column.
    """
    for i, column in enumerate(df.columns):
        max_length = max(len(str(value)) for value in df[column])
        # This is a simple estimation. Adjust the multiplication factor as needed.
        table.columns[i].width = Pt(max_length * 1.5)

def set_cell_borders(cell):
    """
    Set the cell borders to be very narrow on left and right.
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Set left and right border to a specific width
    for border in ["left", "right"]:
        border_elm = OxmlElement(f'w:{border}')
        border_elm.set(qn('w:sz'), '4')  # Width of the border, e.g., 4/8 for a narrow border
        border_elm.set(qn('w:val'), 'single')
        tcPr.append(border_elm)

def add_table_from_dataframe(doc, df, font_name='Arial', font_size=7):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Set header row with font and column widths
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
        set_cell_font(hdr_cells[i], font_name, font_size)
        set_cell_borders(hdr_cells[i])

    set_column_widths(table, df)

    # Add the rest of the data with font
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            set_cell_font(row_cells[i], font_name, font_size)
            set_cell_borders(row_cells[i])

